import pandas as pd
import os

from docx import Document
from docx.oxml.ns import qn
import docx

filename = input('파일 이름(ex.. 면역, 위대한 여정 1부.xls): ')
# filename = '면역, 위대한 여정 1부.xls'

# read
if filename.endswith('.xls'):
    df = pd.read_excel(filename, sheet_name=0, engine='xlrd')
elif filename.endswith('.xlsx'):
    df = pd.read_excel(filename, sheet_name=0, engine='openpyxl')
value = df.values.tolist()
# read

# write
filename_save = filename.split('.')[-2]
docs_dir = os.path.join(os.getcwd(), '{}.docx'.format(filename_save))

doc = Document()
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = docx.shared.Pt(10)
title = doc.add_paragraph()
run = title.add_run(filename_save)
run.bold = True
run.font.size = docx.shared.Pt(16)

for i in value:
    br = doc.add_paragraph()
    tc_in = doc.add_paragraph()
    source = doc.add_paragraph()
    target = doc.add_paragraph()

    br.add_run('')
    tc_in.add_run(i[1])
    source.add_run(i[3])
    target.add_run(i[4])

doc.save(docs_dir)
# write
